from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import tempfile

def generate_cv_doc(cv_data, output_dir="data/output"):
    """Crée un document Word formaté à partir d'un dictionnaire de CV et le sauvegarde dans output_dir."""
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()

    # ======================
    # IDENTITÉ + TITRE
    # ======================
    identite = cv_data.get("identite", {})
    nom = identite.get("nom", "")
    prenom = identite.get("prenom", "")
    age = identite.get("age", "")
    ville = identite.get("ville", "")
    email = identite.get("email", "")
    telephone = identite.get("telephone", "")

    nom_prenom = f"{prenom} {nom}".strip() or "Candidat"

    title = doc.add_heading(f"{nom_prenom}", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Infos personnelles sous le nom
    info_lines = []
    if age: info_lines.append(f"{age} ans")
    if ville: info_lines.append(ville)
    if email: info_lines.append(f"{email}")
    if telephone: info_lines.append(f" {telephone}")

    doc.add_paragraph(" | ".join(info_lines))
    doc.add_paragraph("")  # espace

    # ======================
    # PROFIL
    # ======================
    profil = cv_data.get("profil")
    if profil:
        doc.add_heading("Profil", level=2)
        doc.add_paragraph(profil)

    # ======================
    # EXPÉRIENCES
    # ======================
    doc.add_heading("Expériences professionnelles", level=2)
    experiences = cv_data.get("experiences", [])

    if experiences:
        for exp in experiences:
            poste = exp.get("poste", "")
            organisation = exp.get("organisation", "")
            date_debut = exp.get("date_debut", "")
            date_fin = exp.get("date_fin", "") or "Présent"
            description = exp.get("description", "")

            p = doc.add_paragraph()
            p.add_run(f"• {poste} - {organisation}").bold = True
            doc.add_paragraph(f"🗓 {date_debut} → {date_fin}")
            if description:
                doc.add_paragraph(description, style="List Bullet")
    else:
        doc.add_paragraph("Aucune expérience renseignée.")

    # ======================
    # FORMATIONS
    # ======================
    doc.add_heading("Formations", level=2)
    formations = cv_data.get("formations", [])

    if formations:
        for form in formations:
            diplome = form.get("diplome", "")
            ecole = form.get("ecole", "")
            annee = form.get("annee", "")
            p = doc.add_paragraph()
            p.add_run(f"• {diplome} - {ecole} ({annee})").bold = True
    else:
        doc.add_paragraph("Aucune formation renseignée.")

    # ======================
    # COMPÉTENCES
    # ======================
    doc.add_heading("Compétences", level=2)
    competences = cv_data.get("competences", {})
    tech = competences.get("techniques", [])
    soft = competences.get("soft_skills", [])

    if tech:
        doc.add_paragraph("Compétences techniques : " + ", ".join(tech))
    if soft:
        doc.add_paragraph("Soft skills : " + ", ".join(soft))
    if not tech and not soft:
        doc.add_paragraph("Non renseignées")

    # ======================
    # LANGUES
    # ======================
    doc.add_heading("Langues", level=2)
    langues = cv_data.get("langues", [])
    doc.add_paragraph(", ".join(langues) if langues else "Non renseignées")

    # ======================
    # CENTRES D’INTÉRÊT
    # ======================
    doc.add_heading("Centres d’intérêt", level=2)
    hobbies = cv_data.get("centres_interet", [])
    doc.add_paragraph(", ".join(hobbies) if hobbies else "Non renseignés")

    # ======================
    # EXPORT
    # ======================
    filename = f"{prenom}_{nom}_CV.docx".replace(" ", "_") or "cv_temp.docx"
    file_path = os.path.join(output_dir, filename)
    doc.save(file_path)
    return file_path
